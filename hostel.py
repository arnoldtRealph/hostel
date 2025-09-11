import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime
import pytz
from docx import Document
from docx.shared import Inches
import io
from matplotlib.ticker import MaxNLocator
import logging
from github import Github
import base64
import os

# Configure logging for Streamlit Cloud logs (not UI)
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Set seaborn style for professional charts
sns.set_style("whitegrid")
plt.rcParams['font.size'] = 10
plt.rcParams['axes.titlesize'] = 12
plt.rcParams['axes.labelsize'] = 10
plt.rcParams['xtick.labelsize'] = 9
plt.rcParams['ytick.labelsize'] = 9

# Set page config
st.set_page_config(page_title="Hostel Insident Verslag", layout="wide")

# Custom CSS for futuristic and professional styling
st.markdown("""
    <style>
        /* General layout */
        .stApp {
            background: linear-gradient(135deg, #0f172a 0%, #1e293b 100%);
            font-family: 'Orbitron', sans-serif;
            color: #e2e8f0;
            text-align: center;
        }
        [data-baseweb="baseweb"] {
            background: transparent !important;
        }

        /* Main content */
        .main .block-container {
            padding: 40px;
            background: rgba(15, 23, 42, 0.9);
            border-radius: 20px;
            box-shadow: 0 8px 32px rgba(0, 229, 255, 0.2);
            margin-bottom: 30px;
            transition: transform 0.3s ease, box-shadow 0.3s ease;
            max-width: 1400px;
            margin-left: auto;
            margin-right: auto;
            backdrop-filter: blur(10px);
        }
        .main .block-container:hover {
            transform: translateY(-4px);
            box-shadow: 0 12px 40px rgba(0, 229, 255, 0.3);
        }

        /* Headers */
        h1 {
            color: #00e5ff;
            font-size: 3rem;
            font-weight: 700;
            margin-bottom: 25px;
            text-align: center;
            letter-spacing: 1px;
            text-shadow: 0 0 10px rgba(0, 229, 255, 0.5);
        }
        h2 {
            color: #38bdf8;
            font-size: 2.2rem;
            font-weight: 600;
            margin-top: 30px;
            margin-bottom: 20px;
            border-bottom: 3px solid #0ea5e9;
            padding-bottom: 8px;
            text-align: center;
        }
        h3 {
            color: #38bdf8;
            font-size: 1.6rem;
            font-weight: 500;
            margin-bottom: 15px;
            text-align: center;
        }

        /* Input labels */
        .input-label {
            color: #38bdf8;
            font-size: 1.2rem;
            font-weight: 500;
            margin-bottom: 10px;
            display: block;
            text-align: center;
        }

        /* Buttons */
        .stButton>button, .stDownloadButton>button {
            background: linear-gradient(45deg, #0ea5e9, #38bdf8);
            color: #ffffff !important;
            border: none;
            border-radius: 12px;
            padding: 12px 20px;
            font-size: 1rem;
            font-weight: 500;
            transition: all 0.3s ease;
            max-width: 250px;
            width: 100%;
            margin: 10px auto;
            box-shadow: 0 0 15px rgba(14, 165, 233, 0.5);
            text-align: center;
            display: block;
        }
        .stButton>button:hover, .stDownloadButton>button:hover {
            background: linear-gradient(45deg, #0284c7, #0ea5e9);
            transform: translateY(-2px);
            box-shadow: 0 0 20px rgba(14, 165, 233, 0.7);
        }
        .stButton>button:active, .stDownloadButton>button:active {
            background: linear-gradient(45deg, #0369a1, #0284c7);
            transform: translateY(0);
        }
        .stButton>button:disabled, .stDownloadButton>button:disabled {
            background: #475569;
            color: #94a3b8 !important;
            box-shadow: none;
        }

        /* Delete button */
        .stButton>button.delete-button {
            background: linear-gradient(45deg, #ef4444, #f87171);
            box-shadow: 0 0 15px rgba(239, 68, 68, 0.5);
        }
        .stButton>button.delete-button:hover {
            background: linear-gradient(45deg, #dc2626, #ef4444);
            box-shadow: 0 0 20px rgba(239, 68, 68, 0.7);
        }
        .stButton>button.delete-button:active {
            background: linear-gradient(45deg, #b91c1c, #dc2626);
        }

        /* Download button */
        .stDownloadButton>button {
            background: linear-gradient(45deg, #10b981, #34d399);
        }
        .stDownloadButton>button:hover {
            background: linear-gradient(45deg, #059669, #10b981);
        }
        .stDownloadButton>button:active {
            background: linear-gradient(45deg, #047857, #059669);
        }

        /* Dataframe styling */
        .stDataFrame {
            border: 1px solid #0ea5e9;
            border-radius: 12px;
            overflow-x: auto;
            background: rgba(30, 41, 59, 0.8);
            max-width: 1200px;
            margin: 0 auto;
        }
        .stDataFrame table {
            width: 100%;
            border-collapse: separate;
            border-spacing: 0;
        }
        .stDataFrame th {
            background: #1e293b;
            color: #38bdf8;
            font-weight: 600;
            padding: 15px;
            text-align: left;
            font-size: 1rem;
            border-bottom: 2px solid #0ea5e9;
        }
        .stDataFrame td {
            padding: 15px;
            border-bottom: 1px solid #334155;
            color: #e2e8f0;
            font-size: 0.95rem;
        }
        .stDataFrame tr:nth-child(even) {
            background: rgba(51, 65, 85, 0.3);
        }
        .stDataFrame tr:hover {
            background: rgba(14, 165, 233, 0.2);
        }

        /* Selectbox */
        .stSelectbox {
            background: rgba(30, 41, 59, 0.8);
            border: 1px solid #0ea5e9;
            border-radius: 10px;
            padding: 12px;
            font-size: 1rem;
            transition: border-color 0.3s ease, box-shadow 0.3s ease;
            max-width: 500px;
            width: 100%;
            margin: 10px auto;
            display: block;
            box-shadow: 0 0 10px rgba(14, 165, 233, 0.3);
        }
        .stSelectbox:hover {
            border-color: #38bdf8;
            box-shadow: 0 0 15px rgba(14, 165, 233, 0.5);
        }
        .stSelectbox > div > div {
            min-height: 48px;
            color: #00e5ff !important;
            background: rgba(15, 23, 42, 0.9);
        }
        .stSelectbox [data-baseweb="select"] ul {
            background: #1e293b !important;
            border: 1px solid #0ea5e9 !important;
            border-radius: 8px;
            box-shadow: 0 4px 12px rgba(0, 229, 255, 0.2);
        }
        .stSelectbox [data-baseweb="select"] li {
            color: #e2e8f0 !important;
            padding: 10px;
            transition: background 0.2s ease;
        }
        .stSelectbox [data-baseweb="select"] li:hover {
            background: #0ea5e9 !important;
            color: #ffffff !important;
        }
        .stSelectbox [data-baseweb="select"] li[aria-selected="true"] {
            background: #38bdf8 !important;
            color: #ffffff !important;
            font-weight: 600;
        }

        /* Text area */
        .stTextArea {
            background: rgba(30, 41, 59, 0.8);
            border: 1px solid #0ea5e9;
            border-radius: 10px;
            padding: 12px;
            font-size: 1rem;
            transition: border-color 0.3s ease, box-shadow 0.3s ease;
            max-width: 500px;
            width: 100%;
            margin: 10px auto;
            display: block;
            box-shadow: 0 0 10px rgba(14, 165, 233, 0.3);
            color: #e2e8f0;
        }
        .stTextArea:hover {
            border-color: #38bdf8;
            box-shadow: 0 0 15px rgba(14, 165, 233, 0.5);
        }

        /* Tabs */
        .stTabs {
            display: flex;
            justify-content: center;
            margin-bottom: 20px;
        }
        .stTabs [data-baseweb="tab"] {
            background: #1e293b;
            border-radius: 10px 10px 0 0;
            padding: 12px 24px;
            font-size: 1rem;
            color: #38bdf8;
            margin-right: 6px;
            transition: all 0.3s ease;
        }
        .stTabs [data-baseweb="tab"]:hover {
            background: #334155;
        }
        .stTabs [data-baseweb="tab"][aria-selected="true"] {
            background: #0ea5e9;
            color: #ffffff;
            font-weight: 600;
        }

        /* Charts */
        .stPyplot {
            border-radius: 12px;
            padding: 15px;
            background: rgba(15, 23, 42, 0.9);
            box-shadow: 0 0 15px rgba(14, 165, 233, 0.3);
            max-width: 700px;
            margin: 0 auto;
        }

        /* Separator */
        .custom-divider {
            border-top: 5px solid #0ea5e9;
            margin: 40px auto;
            max-width: 900px;
            border-radius: 2px;
        }

        /* Notification */
        .notification-container > div {
            background: rgba(239, 68, 68, 0.2) !important;
            border: 2px solid #ef4444 !important;
            box-shadow: 0 0 15px rgba(239, 68, 68, 0.3);
        }
        .notification-container > div > h4 {
            color: #f87171 !important;
            text-shadow: 0 0 5px rgba(239, 68, 68, 0.5);
        }
        .notification-container > div > p {
            color: #e2e8f0 !important;
        }

        /* Mobile optimization */
        @media (max-width: 768px) {
            .main .block-container {
                padding: 20px;
                margin-bottom: 20px;
                border-radius: 12px;
            }
            h1 {
                font-size: 2.2rem;
                margin-bottom: 20px;
            }
            h2 {
                font-size: 1.8rem;
                margin-top: 20px;
                margin-bottom: 15px;
            }
            h3 {
                font-size: 1.4rem;
                margin-bottom: 12px;
            }
            .input-label {
                font-size: 1.1rem;
                margin-bottom: 8px;
            }
            .stButton>button, .stDownloadButton>button {
                padding: 10px 16px;
                font-size: 0.95rem;
                max-width: 200px;
                margin: 8px auto;
                border-radius: 8px;
            }
            .stSelectbox, .stTextArea {
                font-size: 0.95rem;
                padding: 10px;
                max-width: 90%;
                margin: 8px auto;
                border-radius: 8px;
            }
            .stDataFrame {
                max-width: 100%;
                font-size: 0.9rem;
            }
            .stTabs [data-baseweb="tab"] {
                padding: 10px 16px;
                font-size: 0.95rem;
                margin-bottom: 5px;
                border-radius: 8px;
                flex: 1 1 auto;
                text-align: center;
            }
        }
    </style>
    <link href="https://fonts.googleapis.com/css2?family=Orbitron:wght@400;500;700&display=swap" rel="stylesheet">
""", unsafe_allow_html=True)

# Load and preprocess learner data
@st.cache_data
def load_learner_data():
    try:
        df = pd.read_csv("learner_list.csv")
        df.columns = df.columns.str.strip()
        df['Learner_Full_Name'] = (df['Leerder van'].fillna('') + ' ' + df['Leerner se naam'].fillna('')).str.strip()
        df = df.rename(columns={
            'BLOK': 'Block',
            'Opvoeder betrokke': 'Teacher',
            'Wat het gebeur': 'Incident',
            'Kategorie': 'Category'
        })
        df['Learner_Full_Name'] = df['Learner_Full_Name'].replace('', 'Onbekend2999')
        df['Block'] = df['Block'].fillna('Onbekend2999')
        df['Teacher'] = df['Teacher'].fillna('Onbekend2999')
        df['Incident'] = df['Incident'].fillna('Onbekend2999')
        df['Category'] = pd.to_numeric(df['Category'], errors='coerce').fillna(1).astype(int).astype(str)
        return df[['Learner_Full_Name', 'Block', 'Teacher', 'Incident', 'Category']]
    except FileNotFoundError:
        logger.warning("learner_list.csv not found. Returning empty DataFrame.")
        return pd.DataFrame(columns=['Learner_Full_Name', 'Block', 'Teacher', 'Incident', 'Category'])

# Load or initialize incident log
def load_incident_log():
    try:
        if os.path.exists("incident_log.csv") and os.path.getsize("incident_log.csv") > 0:
            df = pd.read_csv("incident_log.csv")
        else:
            g = Github(st.secrets["GITHUB_TOKEN"])
            repo = g.get_repo("arnoldtRealph/hostel")
            try:
                contents = repo.get_contents("incident_log.csv", ref="master")
                content = base64.b64decode(contents.content).decode('utf-8')
                df = pd.read_csv(io.StringIO(content))
                df.to_csv("incident_log.csv", index=False)
                logger.info("Incident log fetched from GitHub and saved locally")
            except:
                logger.warning("incident_log.csv does not exist in GitHub repository. Initializing empty DataFrame.")
                df = pd.DataFrame(columns=['Learner_Full_Name', 'Block', 'Teacher', 'Incident', 'Category', 'Comment', 'Date'])

        df.columns = df.columns.str.strip()
        column_mapping = {
            'BLOK': 'Block',
            'Opvoeder betrokke': 'Teacher',
            'Wat het gebeur': 'Incident',
            'Kategorie': 'Category',
            'Leerder Naam': 'Learner_Full_Name',
            'Kommentaar': 'Comment',
            'Datum': 'Date'
        }
        df = df.rename(columns=column_mapping)
        expected_columns = ['Learner_Full_Name', 'Block', 'Teacher', 'Incident', 'Category', 'Comment', 'Date']
        for col in expected_columns:
            if col not in df.columns:
                df[col] = 'Onbekend2999' if col != 'Date' else pd.NaT
        df['Category'] = pd.to_numeric(df['Category'], errors='coerce').fillna(1).astype(int).astype(str)
        df['Date'] = pd.to_datetime(df['Date'], errors='coerce').dt.date
        return df[expected_columns]
    except Exception as e:
        logger.error(f"Error loading incident_log.csv: {e}")
        return pd.DataFrame(columns=['Learner_Full_Name', 'Block', 'Teacher', 'Incident', 'Category', 'Comment', 'Date'])

# Load or initialize general happenings log
def load_happenings_log():
    try:
        if os.path.exists("happenings_log.csv") and os.path.getsize("happenings_log.csv") > 0:
            df = pd.read_csv("happenings_log.csv")
        else:
            g = Github(st.secrets["GITHUB_TOKEN"])
            repo = g.get_repo("arnoldtRealph/hostel")
            try:
                contents = repo.get_contents("happenings_log.csv", ref="master")
                content = base64.b64decode(contents.content).decode('utf-8')
                df = pd.read_csv(io.StringIO(content))
                df.to_csv("happenings_log.csv", index=False)
                logger.info("Happenings log fetched from GitHub and saved locally")
            except:
                logger.warning("happenings_log.csv does not exist in GitHub repository. Initializing empty DataFrame.")
                df = pd.DataFrame(columns=['Learner_Full_Name', 'Block', 'Event', 'Comment', 'Date'])

        df.columns = df.columns.str.strip()
        column_mapping = {
            'Leerder Naam': 'Learner_Full_Name',
            'BLOK': 'Block',
            'Gebeurtenis': 'Event',
            'Kommentaar': 'Comment',
            'Datum': 'Date'
        }
        df = df.rename(columns=column_mapping)
        expected_columns = ['Learner_Full_Name', 'Block', 'Event', 'Comment', 'Date']
        for col in expected_columns:
            if col not in df.columns:
                df[col] = 'Onbekend2999' if col != 'Date' else pd.NaT
        df['Date'] = pd.to_datetime(df['Date'], errors='coerce').dt.date
        return df[expected_columns]
    except Exception as e:
        logger.error(f"Error loading happenings_log.csv: {e}")
        return pd.DataFrame(columns=['Learner_Full_Name', 'Block', 'Event', 'Comment', 'Date'])

# Save incident to log
def save_incident(learner_full_name, block, teacher, incident, category, comment):
    if not all([learner_full_name != 'Kies', block != 'Kies', teacher != 'Kies', incident != 'Kies', category != 'Kies', comment]):
        logger.warning("Incomplete incident fields, saving skipped.")
        st.warning("Alle velde moet ingevul wees om die insident te stoor.")
        return load_incident_log()
        
    incident_log = load_incident_log()
    sa_tz = pytz.timezone('Africa/Johannesburg')
    new_incident = pd.DataFrame({
        'Learner_Full_Name': [learner_full_name],
        'Block': [block],
        'Teacher': [teacher],
        'Incident': [incident],
        'Category': [category],
        'Comment': [comment],
        'Date': [datetime.now(sa_tz).date()]
    })
    incident_log = pd.concat([incident_log, new_incident], ignore_index=True)
    incident_log.to_csv("incident_log.csv", index=False)
    logger.info("Incident saved locally to incident_log.csv")

    try:
        g = Github(st.secrets["GITHUB_TOKEN"])
        repo = g.get_repo("arnoldtRealph/hostel")
        with open("incident_log.csv", "rb") as file:
            content = file.read()
        repo_path = "incident_log.csv"
        try:
            contents = repo.get_contents(repo_path, ref="master")
            repo.update_file(
                path=repo_path,
                message="Updated incident_log.csv with new incident",
                content=content,
                sha=contents.sha,
                branch="master"
            )
            logger.info("Incident log updated on GitHub")
        except:
            repo.create_file(
                path=repo_path,
                message="Created incident_log.csv with new incident",
                content=content,
                branch="master"
            )
            logger.info("Incident log created on GitHub")
    except Exception as e:
        logger.error(f"Failed to push to GitHub: {e}")
        st.error(f"Kon nie na GitHub stoot nie: {e}")

    return incident_log

# Save general happening to log
def save_happening(learner_full_name, block, event, comment):
    if not all([learner_full_name != 'Kies', block != 'Kies', event, comment]):
        logger.warning("Incomplete happening fields, saving skipped.")
        st.warning("Alle velde moet ingevul wees om die gebeurtenis te stoor.")
        return load_happenings_log()
        
    happenings_log = load_happenings_log()
    sa_tz = pytz.timezone('Africa/Johannesburg')
    new_happening = pd.DataFrame({
        'Learner_Full_Name': [learner_full_name],
        'Block': [block],
        'Event': [event],
        'Comment': [comment],
        'Date': [datetime.now(sa_tz).date()]
    })
    happenings_log = pd.concat([happenings_log, new_happening], ignore_index=True)
    happenings_log.to_csv("happenings_log.csv", index=False)
    logger.info("Happening saved locally to happenings_log.csv")

    try:
        g = Github(st.secrets["GITHUB_TOKEN"])
        repo = g.get_repo("arnoldtRealph/hostel")
        with open("happenings_log.csv", "rb") as file:
            content = file.read()
        repo_path = "happenings_log.csv"
        try:
            contents = repo.get_contents(repo_path, ref="master")
            repo.update_file(
                path=repo_path,
                message="Updated happenings_log.csv with new happening",
                content=content,
                sha=contents.sha,
                branch="master"
            )
            logger.info("Happenings log updated on GitHub")
        except:
            repo.create_file(
                path=repo_path,
                message="Created happenings_log.csv with new happening",
                content=content,
                branch="master"
            )
            logger.info("Happenings log created on GitHub")
    except Exception as e:
        logger.error(f"Failed to push to GitHub: {e}")
        st.error(f"Kon nie na GitHub stoot nie: {e}")

    return happenings_log

# Clear a single incident
def clear_incident(index):
    incident_log = load_incident_log()
    if index in incident_log.index:
        incident_log = incident_log.drop(index)
        incident_log.to_csv("incident_log.csv", index=False)
        logger.info(f"Incident at index {index} cleared locally")

        try:
            g = Github(st.secrets["GITHUB_TOKEN"])
            repo = g.get_repo("arnoldtRealph/hostel")
            with open("incident_log.csv", "rb") as file:
                content = file.read()
            repo_path = "incident_log.csv"
            try:
                contents = repo.get_contents(repo_path, ref="master")
                repo.update_file(
                    path=repo_path,
                    message="Updated incident_log.csv after clearing incident",
                    content=content,
                    sha=contents.sha,
                    branch="master"
                )
                logger.info("Incident log updated on GitHub after clearing")
            except:
                repo.create_file(
                    path=repo_path,
                    message="Created incident_log.csv after clearing incident",
                    content=content,
                    branch="master"
                )
                logger.info("Incident log created on GitHub after clearing")
        except Exception as e:
            logger.error(f"Failed to push to GitHub: {e}")
            st.error(f"Kon nie na GitHub stoot nie: {e}")

        return incident_log
    else:
        logger.warning(f"Invalid index {index} for clearing")
        st.warning(f"Ongeldige indeks {index} vir verwydering.")
    return incident_log

# Clear a single happening
def clear_happening(index):
    happenings_log = load_happenings_log()
    if index in happenings_log.index:
        happenings_log = happenings_log.drop(index)
        happenings_log.to_csv("happenings_log.csv", index=False)
        logger.info(f"Happening at index {index} cleared locally")

        try:
            g = Github(st.secrets["GITHUB_TOKEN"])
            repo = g.get_repo("arnoldtRealph/hostel")
            with open("happenings_log.csv", "rb") as file:
                content = file.read()
            repo_path = "happenings_log.csv"
            try:
                contents = repo.get_contents(repo_path, ref="master")
                repo.update_file(
                    path=repo_path,
                    message="Updated happenings_log.csv after clearing happening",
                    content=content,
                    sha=contents.sha,
                    branch="master"
                )
                logger.info("Happenings log updated on GitHub after clearing")
            except:
                repo.create_file(
                    path=repo_path,
                    message="Created happenings_log.csv after clearing happening",
                    content=content,
                    branch="master"
                )
                logger.info("Happenings log created on GitHub after clearing")
        except Exception as e:
            logger.error(f"Failed to push to GitHub: {e}")
            st.error(f"Kon nie na GitHub stoot nie: {e}")

        return happenings_log
    else:
        logger.warning(f"Invalid index {index} for clearing")
        st.warning(f"Ongeldige indeks {index} vir verwydering.")
    return happenings_log

# Generate Word document for incidents and happenings
def generate_word_report(incident_df, happenings_df, learner_name=None):
    doc = Document()
    title = f'Hostel Verslag - {learner_name}' if learner_name else 'Hostel Verslag'
    doc.add_heading(title, 0)

    # Incidents Section
    doc.add_heading('Insident Besonderhede', level=1)
    if not incident_df.empty:
        table = doc.add_table(rows=1, cols=len(incident_df.columns))
        table.style = 'Table Grid'
        for i, col in enumerate(incident_df.columns):
            table.cell(0, i).text = {
                'Learner_Full_Name': 'Leerder Naam',
                'Block': 'Blok',
                'Teacher': 'Toesighouer',
                'Incident': 'Insident',
                'Category': 'Kategorie',
                'Comment': 'Kommentaar',
                'Date': 'Datum'
            }.get(col, col)
        for _, row in incident_df.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(incident_df.columns):
                cells[i].text = str(row[col]) if col != 'Date' else row[col].strftime("%Y-%m-%d") if pd.notnull(row[col]) else 'Onbekend2999'
    else:
        doc.add_paragraph('Geen insidente gerapporteer nie.')

    # General Happenings Section
    doc.add_heading('Algemene Gebeurtenisse', level=1)
    if not happenings_df.empty:
        table = doc.add_table(rows=1, cols=len(happenings_df.columns))
        table.style = 'Table Grid'
        for i, col in enumerate(happenings_df.columns):
            table.cell(0, i).text = {
                'Learner_Full_Name': 'Leerder Naam',
                'Block': 'Blok',
                'Event': 'Gebeurtenis',
                'Comment': 'Kommentaar',
                'Date': 'Datum'
            }.get(col, col)
        for _, row in happenings_df.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(happenings_df.columns):
                cells[i].text = str(row[col]) if col != 'Date' else row[col].strftime("%Y-%m-%d") if pd.notnull(row[col]) else 'Onbekend2999'
    else:
        doc.add_paragraph('Geen algemene gebeurtenisse gerapporteer nie.')

    # Analysis Section
    doc.add_heading('Insident Analise', level=1)

    # Bar chart: Incidents by Category
    if not incident_df.empty:
        fig, ax = plt.subplots(figsize=(4, 2.5))
        category_counts = incident_df['Category'].value_counts().sort_index()
        sns.barplot(x=category_counts.index, y=category_counts.values, ax=ax, palette='Blues')
        ax.set_title('Insidente volgens Kategorie', pad=10, fontsize=12, weight='bold')
        ax.set_xlabel('Kategorie', fontsize=10)
        ax.set_ylabel('Aantal', fontsize=10)
        ax.yaxis.set_major_locator(MaxNLocator(integer=True))
        plt.tight_layout(pad=1.0)
        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', dpi=100, bbox_inches='tight')
        plt.close()
        doc.add_picture(img_stream, width=Inches(3.5))

        # Bar chart: Incidents by Block
        if 'Block' in incident_df.columns:
            fig, ax = plt.subplots(figsize=(4, 2.5))
            block_counts = incident_df['Block'].value_counts()
            sns.barplot(x=block_counts.index, y=block_counts.values, ax=ax, palette='Blues')
            ax.set_title('Insidente volgens Blok', pad=10, fontsize=12, weight='bold')
            ax.set_xlabel('Blok', fontsize=10)
            ax.set_ylabel('Aantal', fontsize=10)
            ax.yaxis.set_major_locator(MaxNLocator(integer=True))
            ax.tick_params(axis='x', rotation=45, labelsize=9)
            plt.tight_layout(pad=1.0)
            img_stream = io.BytesIO()
            plt.savefig(img_stream, format='png', dpi=100, bbox_inches='tight')
            plt.close()
            doc.add_picture(img_stream, width=Inches(3.5))
        else:
            doc.add_paragraph('Geen Blok-data beskikbaar vir analise nie.')
    else:
        doc.add_paragraph('Geen insident-data beskikbaar vir analise nie.')

    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream

# Load data
learner_df = load_learner_data()
incident_log = load_incident_log()
happenings_log = load_happenings_log()

# Main content
with st.container():
    st.title("HOSTEL INSIDENT EN GEBEURTENIS VERSLAG")
    st.subheader("Hoërskool Saul Damon Hostel")

# Initialize session state for sanction notifications
if 'sanction_popups' not in st.session_state:
    st.session_state.sanction_popups = {}

# Compute sanctions based on incident counts
if not incident_log.empty:
    tally_df = incident_log.pivot_table(
        index='Learner_Full_Name',
        columns='Category',
        values='Incident',
        aggfunc='count',
        fill_value=0
    )
    for cat in ['1', '2', '3', '4']:
        if cat not in tally_df.columns:
            tally_df[cat] = 0
    tally_df = tally_df[['1', '2', '3', '4']].reset_index()

    sanctions = []
    for _, row in tally_df.iterrows():
        learner = row['Learner_Full_Name']
        if row['1'] > 5:
            sanctions.append({
                'Learner': learner,
                'Category': '1',
                'Count': int(row['1']),
                'Sanction': 'Waarskuwing en ouerberaad met hosteltoesighouer.'
            })
        if row['2'] > 3:
            sanctions.append({
                'Learner': learner,
                'Category': '2',
                'Count': int(row['2']),
                'Sanction': 'Tydelike verbod op hostelaktiwiteite.'
            })
        if row['3'] > 2:
            sanctions.append({
                'Learner': learner,
                'Category': '3',
                'Count': int(row['3']),
                'Sanction': 'Ouers moet afspraak maak met hostelbestuur.'
            })
        if row['4'] >= 1:
            sanctions.append({
                'Learner': learner,
                'Category': '4',
                'Count': int(row['4']),
                'Sanction': 'Verwysing na dissiplinêre komitee.'
            })

    sanctions_df = pd.DataFrame(sanctions)

    for _, row in sanctions_df.iterrows():
        key = f"{row['Learner']}_{row['Category']}"
        if key not in st.session_state.sanction_popups:
            st.session_state.sanction_popups[key] = True

    st.markdown('<div class="notification-container">', unsafe_allow_html=True)
    any_notifications = False
    for _, row in sanctions_df.iterrows():
        key = f"{row['Learner']}_{row['Category']}"
        if st.session_state.sanction_popups.get(key, False):
            any_notifications = True
            st.markdown(
                f"""
                <div style='padding: 15px; border-radius: 8px;'>
                    <h4 style='margin: 0;'>SANKSIEMELDING</h4>
                    <p style='margin: 5px 0; font-size: 0.9rem;'>
                        Leerder <strong>{row['Learner']}</strong> het {row['Count']} Kategorie {row['Category']} insidente. 
                        Sanksie: {row['Sanction']}
                    </p>
                </div>
                """,
                unsafe_allow_html=True
            )
            if st.button("Opgelos", key=f"sanction_resolve_{key}"):
                st.session_state.sanction_popups[key] = False
                st.rerun()
    if not any_notifications:
        st.markdown(
            """
            <div style='background: rgba(34, 197, 94, 0.2); padding: 15px; border-radius: 8px; border: 2px solid #34d399;'>
                <p style='margin: 0; font-size: 0.9rem;'>Geen aktiewe sanksiemeldings nie.</p>
            </div>
            """,
            unsafe_allow_html=True
        )
    st.markdown('</div>', unsafe_allow_html=True)

# Report new incident
st.header("Rapporteer Nuwe Insident")
with st.container():
    st.markdown('<div class="input-label">Leerder Naam</div>', unsafe_allow_html=True)
    learner_full_name = st.selectbox("", options=['Kies'] + sorted(learner_df['Learner_Full_Name'].unique()), key="learner_full_name")
    
    st.markdown('<div class="input-label">Blok</div>', unsafe_allow_html=True)
    block = st.selectbox("", options=['Kies'] + sorted(learner_df['Block'].unique()), key="block")
    
    st.markdown('<div class="input-label">Toesighouer</div>', unsafe_allow_html=True)
    teacher = st.selectbox("", options=['Kies'] + sorted(learner_df['Teacher'].unique()), key="teacher")
    
    st.markdown('<div class="input-label">Insident</div>', unsafe_allow_html=True)
    incident = st.selectbox("", options=['Kies'] + sorted(learner_df['Incident'].unique()), key="incident")
    
    st.markdown('<div class="input-label">Kategorie</div>', unsafe_allow_html=True)
    category = st.selectbox("", options=['Kies', '1', '2', '3', '4'], key="category")
    
    st.markdown('<div class="input-label">Kommentaar</div>', unsafe_allow_html=True)
    comment = st.text_area("", placeholder="Tik hier...", key="comment")
    
    if st.button("Stoor Insident"):
        incident_log = save_incident(learner_full_name, block, teacher, incident, category, comment)
        st.success("Insident suksesvol gestoor!")

# Report new general happening
st.header("Rapporteer Algemene Gebeurtenis")
with st.container():
    st.markdown('<div class="input-label">Leerder Naam</div>', unsafe_allow_html=True)
    happening_learner = st.selectbox("", options=['Kies'] + sorted(learner_df['Learner_Full_Name'].unique()), key="happening_learner")
    
    st.markdown('<div class="input-label">Blok</div>', unsafe_allow_html=True)
    happening_block = st.selectbox("", options=['Kies'] + sorted(learner_df['Block'].unique()), key="happening_block")
    
    st.markdown('<div class="input-label">Gebeurtenis</div>', unsafe_allow_html=True)
    event = st.text_input("", placeholder="Beskryf die gebeurtenis (bv. Siek, na hostel gestuur)", key="event")
    
    st.markdown('<div class="input-label">Kommentaar</div>', unsafe_allow_html=True)
    happening_comment = st.text_area("", placeholder="Tik hier...", key="happening_comment")
    
    if st.button("Stoor Gebeurtenis"):
        happenings_log = save_happening(happening_learner, happening_block, event, happening_comment)
        st.success("Gebeurtenis suksesvol gestoor!")

# Incident log display
st.header("Insident Log")
if not incident_log.empty:
    display_log = incident_log.copy()
    display_log.index = display_log.index + 1
    st.dataframe(
        display_log,
        use_container_width=True,
        column_config={
            "Learner_Full_Name": st.column_config.TextColumn("Leerder Naam", width="medium"),
            "Block": st.column_config.TextColumn("Blok", width="small"),
            "Teacher": st.column_config.TextColumn("Toesighouer", width="medium"),
            "Incident": st.column_config.TextColumn("Insident", width="medium"),
            "Category": st.column_config.TextColumn("Kategorie", width="small"),
            "Comment": st.column_config.TextColumn("Kommentaar", width="large"),
            "Date": st.column_config.DateColumn("Datum", width="medium", format="YYYY-MM-DD")
        }
    )
    
    st.subheader("Verwyder Insident")
    incident_index = st.number_input("Voer die indeks van die insident in om te verwyder (1-gebaseer)", min_value=1, max_value=len(incident_log), step=1)
    if st.button("Verwyder Insident", key="clear_incident", help="Klik om die geselekteerde insident te verwyder"):
        internal_index = incident_index - 1
        incident_log = clear_incident(internal_index)
        st.success(f"Insident by indeks {incident_index} suksesvol verwyder!")
        st.rerun()
else:
    st.write("Geen insidente in die log nie.")

# General happenings log display
st.header("Algemene Gebeurtenisse Log")
if not happenings_log.empty:
    display_happenings = happenings_log.copy()
    display_happenings.index = display_happenings.index + 1
    st.dataframe(
        display_happenings,
        use_container_width=True,
        column_config={
            "Learner_Full_Name": st.column_config.TextColumn("Leerder Naam", width="medium"),
            "Block": st.column_config.TextColumn("Blok", width="small"),
            "Event": st.column_config.TextColumn("Gebeurtenis", width="medium"),
            "Comment": st.column_config.TextColumn("Kommentaar", width="large"),
            "Date": st.column_config.DateColumn("Datum", width="medium", format="YYYY-MM-DD")
        }
    )
    
    st.subheader("Verwyder Gebeurtenis")
    happening_index = st.number_input("Voer die indeks van die gebeurtenis in om te verwyder (1-gebaseer)", min_value=1, max_value=len(happenings_log), step=1)
    if st.button("Verwyder Gebeurtenis", key="clear_happening", help="Klik om die geselekteerde gebeurtenis te verwyder"):
        internal_index = happening_index - 1
        happenings_log = clear_happening(internal_index)
        st.success(f"Gebeurtenis by indeks {happening_index} suksesvol verwyder!")
        st.rerun()
else:
    st.write("Geen algemene gebeurtenisse in die log nie.")

# Download combined report
st.download_button(
    label="Laai Volledige Verslag af as Word",
    data=generate_word_report(incident_log, happenings_log),
    file_name="hostel_verslag.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

# Filter by learner
st.header("Filter volgens Leerder")
with st.container():
    st.markdown('<div class="input-label">Kies Leerder</div>', unsafe_allow_html=True)
    learner_filter = st.selectbox("", options=['Kies'] + sorted(set(incident_log['Learner_Full_Name'].unique()).union(happenings_log['Learner_Full_Name'].unique())), key="learner_filter")
    
    if learner_filter != 'Kies':
        filtered_incident_log = incident_log[incident_log['Learner_Full_Name'] == learner_filter].copy()
        filtered_happenings_log = happenings_log[happenings_log['Learner_Full_Name'] == learner_filter].copy()
        
        st.subheader(f"Insidente vir {learner_filter}")
        if not filtered_incident_log.empty:
            filtered_incident_log.index = filtered_incident_log.index + 1
            st.dataframe(
                filtered_incident_log,
                use_container_width=True,
                column_config={
                    "Learner_Full_Name": st.column_config.TextColumn("Leerder Naam", width="medium"),
                    "Block": st.column_config.TextColumn("Blok", width="small"),
                    "Teacher": st.column_config.TextColumn("Toesighouer", width="medium"),
                    "Incident": st.column_config.TextColumn("Insident", width="medium"),
                    "Category": st.column_config.TextColumn("Kategorie", width="small"),
                    "Comment": st.column_config.TextColumn("Kommentaar", width="large"),
                    "Date": st.column_config.DateColumn("Datum", width="medium", format="YYYY-MM-DD")
                }
            )
        else:
            st.write("Geen insidente vir hierdie leerder nie.")
        
        st.subheader(f"Algemene Gebeurtenisse vir {learner_filter}")
        if not filtered_happenings_log.empty:
            filtered_happenings_log.index = filtered_happenings_log.index + 1
            st.dataframe(
                filtered_happenings_log,
                use_container_width=True,
                column_config={
                    "Learner_Full_Name": st.column_config.TextColumn("Leerder Naam", width="medium"),
                    "Block": st.column_config.TextColumn("Blok", width="small"),
                    "Event": st.column_config.TextColumn("Gebeurtenis", width="medium"),
                    "Comment": st.column_config.TextColumn("Kommentaar", width="large"),
                    "Date": st.column_config.DateColumn("Datum", width="medium", format="YYYY-MM-DD")
                }
            )
        else:
            st.write("Geen algemene gebeurtenisse vir hierdie leerder nie.")
        
        st.download_button(
            label=f"Laai {learner_filter} se Verslag af",
            data=generate_word_report(filtered_incident_log, filtered_happenings_log, learner_filter),
            file_name=f"verslag_{learner_filter.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="learner_report_download"
        )
    else:
        st.write("Kies 'n leerder om insidente en gebeurtenisse te sien.")

# Today's incidents
st.header("Vandag se Insidente")
today = datetime.now(pytz.timezone('Africa/Johannesburg')).date()
today_incidents = incident_log[incident_log['Date'] == today]
if not today_incidents.empty:
    st.write(f"Totale Insidente Vandag: {len(today_incidents)}")
    fig, ax = plt.subplots(figsize=(4, 2.5))
    category_counts = today_incidents['Category'].value_counts().sort_index()
    sns.barplot(x=category_counts.index, y=category_counts.values, ax=ax, palette='Blues')
    ax.set_title('Insidente volgens Kategorie (Vandag)', pad=10, fontsize=12, weight='bold')
    ax.set_xlabel('Kategorie', fontsize=10)
    ax.set_ylabel('Aantal', fontsize=10)
    ax.yaxis.set_major_locator(MaxNLocator(integer=True))
    plt.tight_layout(pad=1.0)
    st.pyplot(fig)
    plt.close()
else:
    st.write("Geen insidente vandag gerapporteer nie.")

# Today's general happenings
st.header("Vandag se Algemene Gebeurtenisse")
today_happenings = happenings_log[happenings_log['Date'] == today]
if not today_happenings.empty:
    st.write(f"Totale Gebeurtenisse Vandag: {len(today_happenings)}")
    today_happenings_display = today_happenings.copy()
    today_happenings_display.index = today_happenings_display.index + 1
    st.dataframe(
        today_happenings_display,
        use_container_width=True,
        column_config={
            "Learner_Full_Name": st.column_config.TextColumn("Leerder Naam", width="medium"),
            "Block": st.column_config.TextColumn("Blok", width="small"),
            "Event": st.column_config.TextColumn("Gebeurtenis", width="medium"),
            "Comment": st.column_config.TextColumn("Kommentaar", width="large"),
            "Date": st.column_config.DateColumn("Datum", width="medium", format="YYYY-MM-DD")
        }
    )
else:
    st.write("Geen algemene gebeurtenisse vandag gerapporteer nie.")
